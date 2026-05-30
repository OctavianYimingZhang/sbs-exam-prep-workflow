#!/usr/bin/env python3
"""Generate a lecture-first Knowledge Walkthrough DOCX from a plan."""

from __future__ import annotations

import argparse
import json
import re
import shutil
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Cm, Pt, RGBColor  # type: ignore
except Exception as exc:  # pragma: no cover
    raise SystemExit(f"python-docx is required: {exc}")

from knowledge_only_rendering_rules import (
    forbidden_advisory_heading_hits,
    forbidden_advisory_phrase_hits,
    forbidden_non_knowledge_hits,
    repeated_template_label_hits,
)


KNOWLEDGE_WALKTHROUGH_STYLE = {
    "route": "knowledge_walkthrough_docx",
    "margin_cm": 2.0,
    "line_spacing": 1.08,
    "body_alignment": "left",
    "body_font_pt": 10.5,
    "heading_font_pt": 12.0,
    "lecture_heading_font_pt": 14.0,
    "title_font_pt": 14.0,
    "text_color": "black",
    "lecture_page_breaks": True,
}


FORBIDDEN_KEYS = {
    "source_anchor",
    "source_anchors_visible",
    "confidence",
    "evidence",
    "evidence_score",
    "recurrence_count",
    "examiner_operation",
    "discriminator_axis",
    "essay_theme",
    "essay_plan",
    "full_example_essay",
    "practice_question",
    "answer_key",
    "prediction_score",
}

FORBIDDEN_TEXT = {
    "source anchor",
    "confidence",
    "evidence score",
    "recurrence count",
    "examiner operation",
    "discriminator axis",
    "essay plan",
    "full example essay",
    "practice question",
    "answer key",
    "past paper year",
    "prediction score",
    "according to slide",
    "english explanations extracted",
    "ppt page",
    "slide mentions",
    "slides say",
    "the final slide",
    "the first slide",
    "the next slide",
    "the second slide",
    "this slide",
}


GENERIC_SCAFFOLD_HEADINGS = {
    "What This Lecture Is About",
    "What This Module Explains",
    "Knowledge Walkthrough",
    "Key Logic",
    "Knowledge Points",
    "Must Master",
    "Lecture Recap",
}


def safe_filename(text: str, max_len: int = 72) -> str:
    cleaned = re.sub(r"[\/:*?\"<>|]+", "", text)
    cleaned = re.sub(r"\s+", "_", cleaned.strip())
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "", cleaned)
    return cleaned[:max_len].strip("._-") or "knowledge_walkthrough"


def load_plan(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def style_value(plan: dict[str, Any], key: str) -> Any:
    profile = plan.get("route_docx_style_profile")
    if isinstance(profile, dict) and key in profile:
        return profile[key]
    return KNOWLEDGE_WALKTHROUGH_STYLE[key]


def normalize_run(run, size_pt: float | None = None) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size_pt or float(KNOWLEDGE_WALKTHROUGH_STYLE["body_font_pt"]))
    run.font.color.rgb = RGBColor(0, 0, 0)
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def paragraph_size(plan: dict[str, Any], kind: str) -> float:
    profile = plan.get("route_docx_style_profile") or {}
    title_size = style_value(plan, "title_font_pt") if "title_font_pt" in profile else KNOWLEDGE_WALKTHROUGH_STYLE["title_font_pt"]
    return {
        "title": float(title_size),
        "lecture": float(style_value(plan, "lecture_heading_font_pt")),
        "heading": float(style_value(plan, "heading_font_pt")),
        "subheading": float(style_value(plan, "heading_font_pt")),
        "body": float(style_value(plan, "body_font_pt")),
    }[kind]


def normalize_paragraph(paragraph, kind: str, plan: dict[str, Any]) -> None:
    paragraph.paragraph_format.line_spacing = float(style_value(plan, "line_spacing"))
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if kind == "title" else WD_ALIGN_PARAGRAPH.LEFT
    size = paragraph_size(plan, kind)
    for run in paragraph.runs:
        normalize_run(run, size)


def set_document_defaults(doc: Document, plan: dict[str, Any]) -> None:
    margin_cm = float(style_value(plan, "margin_cm"))
    line_spacing = float(style_value(plan, "line_spacing"))
    body_font_pt = float(style_value(plan, "body_font_pt"))
    section = doc.sections[0]
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(body_font_pt)
    doc.styles["Normal"].font.color.rgb = RGBColor(0, 0, 0)
    doc.styles["Normal"].paragraph_format.line_spacing = line_spacing
    doc.styles["Normal"].paragraph_format.space_after = Pt(2)
    for style_name in ["KWTitle", "KWLecture", "KWHeading", "KWSubheading", "KWBody"]:
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, 1)
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(body_font_pt)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = line_spacing
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(2)


def add_marked_text(paragraph, text: str, size_pt: float | None = None) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", str(text))
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        normalize_run(run, size_pt)


def add_paragraph(doc: Document, text: str, plan: dict[str, Any], kind: str = "body"):
    style = {"title": "KWTitle", "lecture": "KWLecture", "heading": "KWHeading", "subheading": "KWSubheading", "body": "KWBody"}[kind]
    paragraph = doc.add_paragraph(style=style)
    add_marked_text(paragraph, text, paragraph_size(plan, kind))
    if kind in {"title", "lecture", "heading", "subheading"}:
        for run in paragraph.runs:
            run.bold = True
    normalize_paragraph(paragraph, kind, plan)
    return paragraph


def add_optional_body(doc: Document, plan: dict[str, Any], value: Any) -> None:
    if value in (None, "", [], {}):
        return
    if isinstance(value, list):
        for item in value:
            text = str(item).strip()
            if text:
                add_paragraph(doc, text, plan, "body")
        return
    text = str(value).strip()
    if text:
        add_paragraph(doc, text, plan, "body")


def add_optional_list(doc: Document, plan: dict[str, Any], values: Any) -> None:
    if not values:
        return
    if not isinstance(values, list):
        add_optional_body(doc, plan, values)
        return
    for item in values:
        text = str(item).strip()
        if text:
            add_paragraph(doc, f"- {text}", plan, "body")


def star_prefix(value: Any) -> str:
    priority = str(value or "").strip()
    return f"{priority} " if priority in {"★★★", "★★", "★"} else ""


def walk_values(value: Any) -> list[tuple[str, Any]]:
    found: list[tuple[str, Any]] = []
    if isinstance(value, dict):
        for key, child in value.items():
            found.append((str(key), child))
            found.extend(walk_values(child))
    elif isinstance(value, list):
        for child in value:
            found.extend(walk_values(child))
    return found


def visible_strings(plan: dict[str, Any]) -> list[str]:
    parts: list[str] = [str(plan.get("title") or "")]
    for lecture in plan.get("lectures", []):
        if not isinstance(lecture, dict):
            continue
        parts.extend(str(lecture.get(key) or "") for key in ["lecture_title", "lecture_overview", "core_logic"])
        for item in lecture.get("module_map", []) or []:
            if isinstance(item, dict):
                parts.extend(str(item.get(key) or "") for key in ["module_title", "one_sentence"])
            else:
                parts.append(str(item))
        for module in lecture.get("modules", []) or []:
            if isinstance(module, dict):
                parts.extend(
                    str(module.get(key) or "")
                    for key in ["module_title", "module_overview", "knowledge_walkthrough", "key_logic"]
                )
                for item in module.get("common_confusions", []) or []:
                    parts.append(str(item))
                for item in module.get("must_master", []) or []:
                    parts.append(str(item))
        for item in lecture.get("lecture_recap", []) or []:
            parts.append(str(item))
    return [part for part in parts if part.strip()]


def validate_plan(plan: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    if not plan.get("lectures"):
        errors.append("walkthrough_requires_lectures")
    if not plan.get("lecture_order"):
        errors.append("walkthrough_requires_lecture_order")
    profile = plan.get("route_docx_style_profile")
    if not isinstance(profile, dict):
        errors.append("walkthrough_requires_route_docx_style_profile")
        profile = {}
    if profile.get("route") != "knowledge_walkthrough_docx":
        errors.append("walkthrough_style_profile_wrong_route")
    margin = profile.get("margin_cm")
    if not isinstance(margin, (int, float)) or abs(float(margin) - 2.0) > 0.2:
        errors.append("walkthrough_style_profile_margin_not_compact")
    spacing = profile.get("line_spacing")
    if not isinstance(spacing, (int, float)) or not (1.0 <= float(spacing) <= 1.2):
        errors.append("walkthrough_style_profile_line_spacing_not_compact")
    if profile.get("body_alignment") != "left":
        errors.append("walkthrough_style_profile_body_alignment_not_left")
    if profile.get("text_color") != "black":
        errors.append("walkthrough_style_profile_text_color_not_black")
    if profile.get("lecture_page_breaks") is not True:
        errors.append("walkthrough_style_profile_missing_lecture_page_breaks")
    for key, value in walk_values(plan):
        if key in FORBIDDEN_KEYS:
            errors.append(f"forbidden_key_visible_in_plan:{key}")
        if isinstance(value, str):
            lowered = value.lower()
            for phrase in FORBIDDEN_TEXT:
                if phrase in lowered:
                    errors.append(f"forbidden_text_in_plan:{phrase}")
            for phrase in forbidden_advisory_phrase_hits(value):
                errors.append(f"forbidden_advisory_phrase_in_plan:{phrase}")
            for heading in forbidden_advisory_heading_hits(value):
                errors.append(f"forbidden_advisory_heading_in_plan:{heading}")
            for category in forbidden_non_knowledge_hits(value):
                errors.append(f"forbidden_non_knowledge_surface_in_plan:{category}")
    combined_visible_text = "\n".join(visible_strings(plan))
    for label in repeated_template_label_hits(combined_visible_text):
        errors.append(f"repeated_rigid_template_label:{label}")
    for heading in GENERIC_SCAFFOLD_HEADINGS:
        if re.search(rf"(?im)^\s*{re.escape(heading)}\s*:?\s*$", combined_visible_text):
            errors.append(f"generic_scaffold_heading_visible:{heading}")
    for lecture in plan.get("lectures", []):
        if not lecture.get("modules"):
            errors.append(f"lecture_missing_modules:{lecture.get('lecture_id', 'unknown')}")
    return sorted(set(errors))


def write_docx(plan: dict[str, Any], output_dir: Path, qa_dir: Path, strict: bool) -> dict[str, Any]:
    errors = validate_plan(plan)
    if strict and errors:
        return {"status": "fail", "qa_flags": errors, "documents": []}

    doc = Document()
    set_document_defaults(doc, plan)
    title = plan.get("title") or "Lecture Knowledge Walkthrough"
    add_paragraph(doc, str(title), plan, "title")

    manifest = {
        "walkthrough_id": plan.get("walkthrough_id"),
        "target_group_key": plan.get("target_group_key"),
        "title": title,
        "lectures": [],
        "qa_flags": errors,
    }

    for lecture_index, lecture in enumerate(plan.get("lectures", [])):
        if lecture_index > 0 and style_value(plan, "lecture_page_breaks"):
            doc.add_page_break()
        lecture_title = lecture.get("lecture_title", "Lecture")
        add_paragraph(doc, f"Lecture: {lecture_title}", plan, "lecture")
        add_optional_body(doc, plan, lecture.get("lecture_overview"))
        add_optional_body(doc, plan, lecture.get("core_logic"))

        module_map = lecture.get("module_map") or []
        if module_map:
            add_paragraph(doc, "Knowledge map", plan, "subheading")
            for item in module_map:
                if isinstance(item, dict):
                    title_text = str(item.get("module_title") or "Knowledge area").strip()
                    summary = str(item.get("one_sentence") or "").strip()
                    add_paragraph(doc, f"{title_text}: {summary}" if summary else title_text, plan, "body")
                else:
                    add_paragraph(doc, str(item), plan, "body")

        lecture_manifest = {"lecture_id": lecture.get("lecture_id"), "lecture_title": lecture_title, "modules": []}
        for module in lecture.get("modules", []):
            module_title = str(module.get("module_title") or "Knowledge module").strip()
            add_paragraph(doc, f"{star_prefix(module.get('priority'))}{module_title}", plan, "subheading")
            add_optional_body(doc, plan, module.get("module_overview"))
            add_optional_body(doc, plan, module.get("knowledge_walkthrough"))
            add_optional_body(doc, plan, module.get("key_logic"))
            common_confusions = module.get("common_confusions", [])
            if common_confusions:
                add_paragraph(doc, "Key distinctions", plan, "subheading")
                add_optional_list(doc, plan, common_confusions)
            must_master = module.get("must_master", [])
            if must_master:
                add_paragraph(doc, "Core knowledge points", plan, "subheading")
                add_optional_list(doc, plan, must_master)
            lecture_manifest["modules"].append({"module_id": module.get("module_id"), "module_title": module.get("module_title")})

        lecture_recap = lecture.get("lecture_recap", [])
        if lecture_recap:
            add_paragraph(doc, "Synthesis", plan, "subheading")
            add_optional_list(doc, plan, lecture_recap)
        manifest["lectures"].append(lecture_manifest)

    output_dir.mkdir(parents=True, exist_ok=True)
    qa_dir.mkdir(parents=True, exist_ok=True)
    filename = "Lecture_Knowledge_Walkthrough.docx"
    docx_path = output_dir / filename
    doc.save(docx_path)
    manifest["documents"] = [{"docx_path": str(docx_path), "filename": filename}]
    manifest["route_docx_style_profile"] = plan.get("route_docx_style_profile") or KNOWLEDGE_WALKTHROUGH_STYLE
    manifest["status"] = "pass" if not errors else "warn"
    manifest_path = qa_dir / "knowledge_walkthrough_manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return manifest


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--plan", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--qa-dir", type=Path)
    parser.add_argument("--clean", action="store_true")
    parser.add_argument("--strict", action="store_true")
    parser.add_argument("--deliverable-only", action="store_true")
    args = parser.parse_args()

    output_dir = args.output_dir
    qa_dir = args.qa_dir or (output_dir if not args.deliverable_only else output_dir.parent / "knowledge_walkthrough_internal_qa")
    if args.clean:
        if output_dir.exists():
            shutil.rmtree(output_dir)
        if qa_dir.exists() and qa_dir != output_dir:
            shutil.rmtree(qa_dir)

    result = write_docx(load_plan(args.plan), output_dir, qa_dir, args.strict)
    print(json.dumps(result, indent=2))
    return 0 if result.get("status") in {"pass", "warn"} else 1


if __name__ == "__main__":
    raise SystemExit(main())
