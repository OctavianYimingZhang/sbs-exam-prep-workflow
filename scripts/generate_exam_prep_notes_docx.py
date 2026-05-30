#!/usr/bin/env python3
"""Generate Academic Exam-Ready Notes DOCX from an ExamPrepNotesPlan."""

from __future__ import annotations

import argparse
import json
import re
import shutil
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Cm, Pt, RGBColor  # type: ignore
except Exception as exc:  # pragma: no cover
    raise SystemExit(f"python-docx is required: {exc}")

from knowledge_only_rendering_rules import (
    forbidden_advisory_heading_hits,
    forbidden_advisory_phrase_hits,
    forbidden_non_knowledge_hits,
    repeated_template_label_hits,
)


FORBIDDEN_PUBLIC_PHRASES = [
    "Course-Level Exam Map",
    "assessment includes",
    "summative examination",
    "Section A asks",
    "Section B is",
    "Coverage note:",
    "no mark scheme",
    "historical papers",
    "older papers",
    "current regime",
    "ELM should be checked",
    "source coverage",
    "extraction quality",
]

FORBIDDEN_INTERNAL_HEADINGS = {
    "Exam Specificity",
    "Core Exam Claim",
    "Exam Use",
    "Common Error / Trap",
    "Must Master",
    "How To Answer This Exam",
}

EXAM_PREP_NOTES_STYLE = {
    "margin_cm": 2.0,
    "line_spacing": 1.08,
    "body_alignment": "left",
    "body_font_pt": 10.5,
    "heading_font_pt": 12.0,
    "lecture_heading_font_pt": 14.0,
    "text_color": "black",
}

BLOCK_LABELS = {
    "definitions": "Definitions",
    "key_points": "Key Points",
    "criteria": "Criteria",
    "steps": "Steps",
    "mechanism": "Mechanism",
    "equation": "Equation",
    "calculation_logic": "Calculation Logic",
    "graph_logic": "Graph Logic",
    "comparison": "Comparison",
    "example": "Example",
    "limitation": "Limitation",
    "worked_example": "Worked Example",
    "diagnostic_pattern": "Diagnostic Pattern",
    "control": "Control",
}

ESSENTIAL_LABEL_TYPES = {"equation", "worked_example", "diagnostic_pattern", "control", "comparison"}
ESSENTIAL_LABEL_TEXT = {"equation", "worked example", "diagnostic pattern", "control", "comparison", "table"}


def load_plan(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def style_value(plan: dict[str, Any], key: str) -> Any:
    profile = plan.get("route_docx_style_profile")
    if isinstance(profile, dict) and key in profile:
        return profile[key]
    return EXAM_PREP_NOTES_STYLE[key]


def normalize_run(run, size_pt: float | None = None) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size_pt or EXAM_PREP_NOTES_STYLE["body_font_pt"])
    run.font.color.rgb = RGBColor(0, 0, 0)
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def normalize_paragraph(paragraph, kind: str, plan: dict[str, Any]) -> None:
    paragraph.paragraph_format.line_spacing = float(style_value(plan, "line_spacing"))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if kind == "title" else WD_ALIGN_PARAGRAPH.LEFT
    size = {
        "title": 14.0,
        "lecture": float(style_value(plan, "lecture_heading_font_pt")),
        "heading": float(style_value(plan, "heading_font_pt")),
        "body": float(style_value(plan, "body_font_pt")),
    }[kind]
    for run in paragraph.runs:
        normalize_run(run, size)


def set_document_defaults(doc: Document, plan: dict[str, Any]) -> None:
    margin_cm = float(style_value(plan, "margin_cm"))
    line_spacing = float(style_value(plan, "line_spacing"))
    body_font_pt = float(style_value(plan, "body_font_pt"))
    section = doc.sections[0]
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)
    for style_name in ["Normal", "EPNTitle", "EPNLecture", "EPNHeading", "EPNBody"]:
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, 1)
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(body_font_pt)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = line_spacing


def add_marked_text(paragraph, text: str, size_pt: float | None = None) -> None:
    for part in re.split(r"(\*\*[^*]+\*\*)", str(text)):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        normalize_run(run, size_pt)


def add_paragraph(doc: Document, text: str, plan: dict[str, Any], kind: str = "body"):
    style = {"title": "EPNTitle", "lecture": "EPNLecture", "heading": "EPNHeading", "body": "EPNBody"}[kind]
    paragraph = doc.add_paragraph(style=style)
    size = {
        "title": 14.0,
        "lecture": float(style_value(plan, "lecture_heading_font_pt")),
        "heading": float(style_value(plan, "heading_font_pt")),
        "body": float(style_value(plan, "body_font_pt")),
    }[kind]
    add_marked_text(paragraph, text, size)
    if kind in {"title", "lecture", "heading"}:
        for run in paragraph.runs:
            run.bold = True
    normalize_paragraph(paragraph, kind, plan)
    return paragraph


def label_policy_mode(plan: dict[str, Any]) -> str:
    for key in ["knowledge_surface_contract", "surface_contract"]:
        contract = plan.get(key)
        if isinstance(contract, dict):
            policy = contract.get("label_policy")
            if isinstance(policy, dict) and policy.get("mode"):
                return str(policy["mode"])
    language_profile = plan.get("output_language_profile")
    if isinstance(language_profile, dict):
        policies = set(str(item) for item in language_profile.get("public_label_policy", []))
        if "explicit_user_requested" in policies:
            return "explicit_user_requested"
        if "suppress_nonessential_labels" in policies or "semantic_sparse" in policies:
            return "semantic_sparse"
    return "semantic_sparse"


def public_labels_enabled(plan: dict[str, Any]) -> bool:
    return label_policy_mode(plan) != "semantic_sparse"


def label_is_semantically_essential(block: dict[str, Any], label: str | None) -> bool:
    if block.get("force_label") is True or block.get("label_decision") == "keep":
        return True
    block_type = str(block.get("block_type") or "").strip().casefold()
    if block_type in ESSENTIAL_LABEL_TYPES:
        return True
    normalized_label = str(label or "").strip().casefold()
    return normalized_label in ESSENTIAL_LABEL_TEXT


def block_label(block: dict[str, Any], plan: dict[str, Any]) -> str | None:
    explicit = block.get("label")
    label = str(explicit).strip().rstrip(":") if explicit is not None and str(explicit).strip() else None
    block_type = str(block.get("block_type") or "")
    if label_policy_mode(plan) == "semantic_sparse" and not label_is_semantically_essential(block, label):
        return None
    if label:
        return label
    if not public_labels_enabled(plan) and block_type not in ESSENTIAL_LABEL_TYPES:
        return None
    requested_language = str((plan.get("output_language_profile") or {}).get("requested_language") or "").strip().lower()
    if requested_language and requested_language not in {"english", "en", "en-gb", "en-us"}:
        return None
    generated_label = BLOCK_LABELS.get(block_type)
    return generated_label if label_is_semantically_essential(block, generated_label) or public_labels_enabled(plan) else None


def add_label_block(doc: Document, plan: dict[str, Any], label: str | None, value: Any) -> None:
    if value in (None, "", [], {}):
        return
    if label:
        add_paragraph(doc, f"{label}:", plan, "heading")
    if isinstance(value, list):
        for item in value:
            add_paragraph(doc, f"- {item}", plan, "body")
    else:
        add_paragraph(doc, str(value), plan, "body")


def point_kind_from_card(card: dict[str, Any]) -> str:
    mapping = {
        "definition": "definition",
        "mechanism": "mechanism",
        "criteria list": "criteria_list",
        "comparison": "comparison",
        "calculation": "calculation",
        "graph interpretation": "graph_or_data_interpretation",
        "method workflow": "method_workflow",
        "case justification": "compact_background",
        "background": "compact_background",
    }
    return mapping.get(str(card.get("exam_specificity") or "").strip().casefold(), "compact_background")


def derive_public_points(plan: dict[str, Any]) -> list[dict[str, Any]]:
    points: list[dict[str, Any]] = []
    default_lecture = ""
    if isinstance(plan.get("lecture_mapping"), list) and plan["lecture_mapping"]:
        first_lecture = plan["lecture_mapping"][0]
        if isinstance(first_lecture, dict):
            default_lecture = str(first_lecture.get("lecture_session_id") or "")
    for card in plan.get("knowledge_cards", []):
        if not isinstance(card, dict) or not card.get("student_visible", True):
            continue
        blocks: list[dict[str, Any]] = []
        block_specs = [
            ("key_definitions", "definitions"),
            ("criteria_components_steps", "criteria"),
            ("mechanism_or_process_logic", "mechanism"),
            ("canonical_example", "example"),
        ]
        for source_field, block_type in block_specs:
            value = card.get(source_field)
            if value:
                blocks.append({"block_type": block_type, "label": None, "content": value, "covered_atomic_units": []})
        main_text = " ".join(
            str(card.get(field) or "").strip()
            for field in ["core_exam_claim", "exam_ready_knowledge_synthesis"]
            if str(card.get(field) or "").strip()
        )
        covered = [str(item) for item in card.get("protected_source_items_covered", []) if str(item).strip()]
        points.append(
            {
                "point_id": f"public_{card.get('card_id', len(points) + 1)}",
                "lecture_session_id": str(card.get("lecture_session_id") or default_lecture or card.get("section_id") or "lecture_001"),
                "source_card_ids": [str(card.get("card_id"))] if card.get("card_id") else [],
                "point_title": str(card.get("module_title") or "Knowledge point"),
                "priority": card.get("priority") or "★",
                "point_kind": point_kind_from_card(card),
                "main_text": main_text or str(card.get("exam_ready_knowledge_synthesis") or ""),
                "blocks": blocks,
                "covered_atomic_units": covered or [str(card.get("kp_id") or card.get("card_id") or "covered_item")],
            }
        )
    return points


def visible_text_from_plan(plan: dict[str, Any]) -> str:
    parts: list[str] = []
    for key in ["title", "course_knowledge_map"]:
        if plan.get(key):
            parts.append(str(plan[key]))
    for section in plan.get("course_sections", []):
        parts.extend(str(section.get(key, "")) for key in ["section_title", "section_function"])
    for point in plan.get("public_output_points", []) or []:
        if not isinstance(point, dict):
            continue
        parts.extend(str(point.get(key, "")) for key in ["point_title", "priority", "main_text"])
        for block in point.get("blocks", []):
            if isinstance(block, dict):
                parts.append(str(block.get("label") or ""))
                parts.append(json.dumps(block.get("content", ""), ensure_ascii=False))
    return "\n".join(parts)


def validate_plan(plan: dict[str, Any]) -> list[str]:
    failures: list[str] = []
    if plan.get("object_type") != "ExamPrepNotesPlan":
        failures.append("plan_object_type_not_exam_prep_notes_plan")
    cards = [card for card in plan.get("knowledge_cards", []) if card.get("student_visible", True)]
    points = [point for point in plan.get("public_output_points", []) if isinstance(point, dict)]
    if not points:
        failures.append("no_public_output_points")
    if not plan.get("knowledge_only_student_view_binding"):
        failures.append("missing_knowledge_only_student_view_binding")
    for card in cards:
        for field in ["module_title", "priority", "exam_specificity", "core_exam_claim", "exam_ready_knowledge_synthesis"]:
            if not card.get(field):
                failures.append(f"visible_card_missing_{field}:{card.get('card_id', 'unknown')}")
        if card.get("priority") not in {"★★★", "★★", "★"}:
            failures.append(f"visible_card_bad_priority:{card.get('card_id', 'unknown')}")
    for point in points:
        for field in ["lecture_session_id", "point_title", "priority", "point_kind", "main_text", "covered_atomic_units"]:
            if not point.get(field):
                failures.append(f"public_point_missing_{field}:{point.get('point_id', 'unknown')}")
        if point.get("priority") not in {"★★★", "★★", "★"}:
            failures.append(f"public_point_bad_priority:{point.get('point_id', 'unknown')}")
    visible_text = visible_text_from_plan(plan)
    for phrase in FORBIDDEN_PUBLIC_PHRASES:
        if phrase.lower() in visible_text.lower():
            failures.append(f"forbidden_public_phrase:{phrase}")
    for heading in FORBIDDEN_INTERNAL_HEADINGS:
        if re.search(rf"(?im)^\s*{re.escape(heading)}\s*:?\s*$", visible_text):
            failures.append(f"forbidden_internal_heading:{heading}")
    for phrase in forbidden_advisory_phrase_hits(visible_text):
        failures.append(f"forbidden_advisory_phrase:{phrase}")
    for heading in forbidden_advisory_heading_hits(visible_text):
        failures.append(f"forbidden_advisory_heading:{heading}")
    for category in forbidden_non_knowledge_hits(visible_text):
        failures.append(f"forbidden_non_knowledge_surface:{category}")
    for label in repeated_template_label_hits(visible_text):
        failures.append(f"repeated_rigid_template_label:{label}")
    return sorted(set(failures))


def write_docx(plan: dict[str, Any], output_dir: Path, qa_dir: Path, strict: bool) -> dict[str, Any]:
    failures = validate_plan(plan)
    if strict and failures:
        return {"status": "fail", "qa_flags": failures, "documents": []}

    doc = Document()
    set_document_defaults(doc, plan)
    title = str(plan.get("title") or "Academic Exam-Ready Knowledge Notes")
    add_paragraph(doc, title, plan, "title")

    add_paragraph(doc, "Course Knowledge Map", plan, "heading")
    course_map = plan.get("course_knowledge_map")
    if course_map:
        add_paragraph(doc, str(course_map), plan, "body")
    else:
        add_paragraph(doc, "Core topics are grouped by concepts, mechanisms, method workflows, calculations and interpretation rules.", plan, "body")

    lectures = [lecture for lecture in plan.get("lecture_mapping", []) if isinstance(lecture, dict)]
    lecture_by_id = {str(lecture.get("lecture_session_id")): lecture for lecture in lectures}
    lecture_order = [str(lecture.get("lecture_session_id")) for lecture in lectures]
    points = [point for point in plan.get("public_output_points", []) if isinstance(point, dict)] or derive_public_points(plan)
    grouped: dict[str, list[dict[str, Any]]] = {lecture_id: [] for lecture_id in lecture_order}
    for point in points:
        lecture_id = str(point.get("lecture_session_id") or "")
        if lecture_id not in grouped:
            grouped[lecture_id] = []
            lecture_order.append(lecture_id)
            lecture_by_id[lecture_id] = {"title": lecture_id or "Lecture"}
        grouped[lecture_id].append(point)

    for lecture_id in lecture_order:
        lecture_points = grouped.get(lecture_id, [])
        if not lecture_points:
            continue
        doc.add_page_break()
        lecture = lecture_by_id.get(lecture_id, {})
        lecture_title = str(lecture.get("title") or lecture_id or "Lecture")
        add_paragraph(doc, f"Lecture: {lecture_title}", plan, "lecture")
        for point in lecture_points:
            title_text = str(point.get("point_title") or "Knowledge Point").strip()
            priority = str(point.get("priority") or "★").strip()
            add_paragraph(doc, f"{priority} {title_text}", plan, "heading")
            if point.get("main_text"):
                add_paragraph(doc, str(point["main_text"]), plan, "body")
            for block in point.get("blocks", []):
                if not isinstance(block, dict):
                    continue
                add_label_block(doc, plan, block_label(block, plan), block.get("content"))

    output_dir.mkdir(parents=True, exist_ok=True)
    qa_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / "Lecture_Knowledge_Walkthrough.docx"
    doc.save(docx_path)
    manifest = {
        "status": "pass" if not failures else "warn",
        "notes_plan_id": plan.get("notes_plan_id"),
        "target_group_key": plan.get("target_group_key"),
        "qa_flags": failures,
        "documents": [{"docx_path": str(docx_path), "filename": docx_path.name}],
    }
    (qa_dir / "exam_prep_notes_manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return manifest


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--plan", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--qa-dir", type=Path)
    parser.add_argument("--clean", action="store_true")
    parser.add_argument("--strict", action="store_true")
    parser.add_argument("--deliverable-only", action="store_true")
    args = parser.parse_args()

    output_dir = args.output_dir
    qa_dir = args.qa_dir or (output_dir if not args.deliverable_only else output_dir.parent / "exam_prep_notes_internal_qa")
    if args.clean:
        if output_dir.exists():
            shutil.rmtree(output_dir)
        if qa_dir.exists() and qa_dir != output_dir:
            shutil.rmtree(qa_dir)

    result = write_docx(load_plan(args.plan), output_dir, qa_dir, args.strict)
    print(json.dumps(result, indent=2))
    return 0 if result.get("status") in {"pass", "warn"} else 1


if __name__ == "__main__":
    raise SystemExit(main())
