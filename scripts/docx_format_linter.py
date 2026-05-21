#!/usr/bin/env python3
"""Validate Example Essay DOCX formatting, highlighting, and source maps."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX  # type: ignore
    from docx.shared import Cm, Pt  # type: ignore
except Exception as exc:  # pragma: no cover
    raise SystemExit(f"python-docx is required: {exc}")


AUTHOR_YEAR_RE = re.compile(r"\([A-Z][A-Za-z'’-]+(?:\s+et\s+al\.|\s+and\s+[A-Z][A-Za-z'’-]+)?(?:,\s*|\s+)(?:19|20)\d{2}[a-z]?\)")
WORD_RE = re.compile(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?")
CM_TOL = 0.08
PT_TOL = 0.25


def words(text: str) -> list[str]:
    return WORD_RE.findall(text)


def emu_to_cm(value) -> float:
    return float(value.cm)


def pt_value(value) -> float | None:
    if value is None:
        return None
    return float(value.pt)


def close(actual: float | None, expected: float, tol: float) -> bool:
    return actual is not None and abs(actual - expected) <= tol


def load_source_map(docx_path: Path, explicit_map: Path | None) -> dict[str, Any] | None:
    if explicit_map:
        return json.loads(explicit_map.read_text(encoding="utf-8"))
    candidate = docx_path.with_name(docx_path.stem.split("_")[0] + "_source_map.json")
    if candidate.exists():
        return json.loads(candidate.read_text(encoding="utf-8"))
    candidate = docx_path.with_name(docx_path.stem + "_source_map.json")
    if candidate.exists():
        return json.loads(candidate.read_text(encoding="utf-8"))
    return None


def run_font_name(run) -> str | None:
    return run.font.name


def run_font_size_pt(run) -> float | None:
    return pt_value(run.font.size)


def paragraph_kind_from_map(source_map: dict[str, Any] | None, index: int) -> str:
    if not source_map:
        if index == 1:
            return "title"
        return "body"
    paragraphs = source_map.get("paragraphs", [])
    if index - 1 < len(paragraphs):
        return paragraphs[index - 1].get("kind", "body")
    return "body"


def lint_docx(docx_path: Path, source_map_path: Path | None = None) -> dict[str, Any]:
    doc = Document(docx_path)
    source_map = load_source_map(docx_path, source_map_path)
    failures: list[dict[str, Any]] = []
    warnings: list[dict[str, Any]] = []

    section = doc.sections[0]
    format_checks = {
        "a4": close(emu_to_cm(section.page_width), 21.0, CM_TOL) and close(emu_to_cm(section.page_height), 29.7, CM_TOL),
        "margins_2_5_cm": all(
            close(emu_to_cm(value), 2.5, CM_TOL)
            for value in [section.top_margin, section.bottom_margin, section.left_margin, section.right_margin]
        ),
        "arial_10": True,
        "line_spacing_1_5": True,
        "paragraph_spacing_zero": True,
        "body_justified": True,
        "title_centered": True,
        "subtitles_left": True,
        "no_empty_spacer_paragraphs": True,
    }
    if not format_checks["a4"]:
        failures.append({"type": "page_size_not_a4"})
    if not format_checks["margins_2_5_cm"]:
        failures.append({"type": "margins_not_2_5_cm"})

    total_body_words = 0
    yellow_words = 0
    green_words = 0
    paragraphs_missing_lecture_anchor = []
    green_runs_missing_citation = []
    green_runs_missing_read_source = []
    yellow_runs_missing_extra_reading_anchor = []
    unexpected_highlights = []

    source_paragraphs = source_map.get("paragraphs", []) if source_map else []
    allowed_author_years = set()
    if source_map:
        for paragraph in source_paragraphs:
            for run in paragraph.get("runs", []):
                citation = run.get("in_text_citation")
                if citation:
                    allowed_author_years.add(citation)

    visible_paragraph_index = 0
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            format_checks["no_empty_spacer_paragraphs"] = False
            failures.append({"type": "empty_spacer_paragraph"})
            continue
        visible_paragraph_index += 1
        kind = paragraph_kind_from_map(source_map, visible_paragraph_index)
        source_para = source_paragraphs[visible_paragraph_index - 1] if visible_paragraph_index - 1 < len(source_paragraphs) else {}

        before = pt_value(paragraph.paragraph_format.space_before)
        after = pt_value(paragraph.paragraph_format.space_after)
        if not close(before if before is not None else 0.0, 0.0, PT_TOL) or not close(after if after is not None else 0.0, 0.0, PT_TOL):
            format_checks["paragraph_spacing_zero"] = False
            failures.append({"type": "paragraph_spacing_not_zero", "paragraph": visible_paragraph_index})
        if paragraph.paragraph_format.line_spacing != 1.5:
            format_checks["line_spacing_1_5"] = False
            failures.append({"type": "line_spacing_not_1_5", "paragraph": visible_paragraph_index, "line_spacing": paragraph.paragraph_format.line_spacing})

        if kind == "title" and paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            format_checks["title_centered"] = False
            failures.append({"type": "title_not_centered", "paragraph": visible_paragraph_index})
        elif kind in {"subtitle", "heading"} and paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT:
            format_checks["subtitles_left"] = False
            failures.append({"type": "subtitle_or_heading_not_left", "paragraph": visible_paragraph_index})
        elif kind == "body" and paragraph.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            format_checks["body_justified"] = False
            failures.append({"type": "body_not_justified", "paragraph": visible_paragraph_index})

        if kind == "body":
            if source_map and not source_para.get("lecture_anchors"):
                paragraphs_missing_lecture_anchor.append(visible_paragraph_index)

        source_runs = source_para.get("runs", []) if source_para else []
        visible_runs = [run for run in paragraph.runs if run.text]
        for r_idx, run in enumerate(visible_runs, start=1):
            if run_font_name(run) != "Arial":
                format_checks["arial_10"] = False
                failures.append({"type": "run_font_not_arial", "paragraph": visible_paragraph_index, "run": r_idx, "font": run_font_name(run)})
            size = run_font_size_pt(run)
            if not close(size, 10.0, PT_TOL):
                format_checks["arial_10"] = False
                failures.append({"type": "run_font_size_not_10", "paragraph": visible_paragraph_index, "run": r_idx, "size": size})

            source_run = source_runs[r_idx - 1] if r_idx - 1 < len(source_runs) else {}
            wc = len(words(run.text))
            if kind == "body":
                total_body_words += wc
            highlight = run.font.highlight_color
            if highlight == WD_COLOR_INDEX.YELLOW:
                yellow_words += wc
                if source_run.get("source_type") != "extra_reading_book" or not source_run.get("source_anchor"):
                    yellow_runs_missing_extra_reading_anchor.append({"paragraph": visible_paragraph_index, "run": r_idx, "text": run.text[:120]})
            elif highlight == WD_COLOR_INDEX.BRIGHT_GREEN:
                green_words += wc
                citation = source_run.get("in_text_citation")
                if not citation and not AUTHOR_YEAR_RE.search(run.text):
                    green_runs_missing_citation.append({"paragraph": visible_paragraph_index, "run": r_idx, "text": run.text[:120]})
                if source_run.get("source_type") != "citation_original_source" or source_run.get("citation_original_read") is not True:
                    green_runs_missing_read_source.append({"paragraph": visible_paragraph_index, "run": r_idx, "text": run.text[:120]})
            elif highlight is not None:
                unexpected_highlights.append({"paragraph": visible_paragraph_index, "run": r_idx, "highlight": str(highlight)})

    if paragraphs_missing_lecture_anchor:
        failures.append({"type": "paragraphs_missing_lecture_anchor", "paragraphs": paragraphs_missing_lecture_anchor})
    if green_runs_missing_citation:
        failures.append({"type": "green_highlight_missing_citation", "runs": green_runs_missing_citation})
    if green_runs_missing_read_source:
        failures.append({"type": "green_highlight_missing_read_original_source", "runs": green_runs_missing_read_source})
    if yellow_runs_missing_extra_reading_anchor:
        failures.append({"type": "yellow_highlight_missing_extra_reading_anchor", "runs": yellow_runs_missing_extra_reading_anchor})
    if unexpected_highlights:
        warnings.append({"type": "unexpected_highlight_colour", "runs": unexpected_highlights})

    extra_reading_ratio = yellow_words / total_body_words if total_body_words else 0.0
    extra_status = (source_map or {}).get("extra_reading_status")
    if yellow_words and extra_reading_ratio > 0.15:
        failures.append({"type": "extra_reading_ratio_above_15_percent", "ratio": extra_reading_ratio})
    if extra_status in {"supplied_found", "found", "chapter_found"} and extra_reading_ratio < 0.10:
        failures.append({"type": "extra_reading_ratio_below_10_percent", "ratio": extra_reading_ratio})

    used_citations = set(AUTHOR_YEAR_RE.findall("\n".join(p.text for p in doc.paragraphs)))
    if source_map and allowed_author_years and not used_citations.issubset(allowed_author_years):
        failures.append({"type": "essay_contains_author_year_not_in_source_map", "citations": sorted(used_citations - allowed_author_years)})

    report = {
        "docx": str(docx_path),
        "status": "pass" if not failures else "fail",
        "format": format_checks,
        "source_grounding": {
            "paragraphs_missing_lecture_anchor": paragraphs_missing_lecture_anchor,
            "green_runs_missing_citation": green_runs_missing_citation,
            "yellow_runs_missing_extra_reading_anchor": yellow_runs_missing_extra_reading_anchor,
            "green_runs_missing_read_original_source": green_runs_missing_read_source,
            "extra_reading_ratio": round(extra_reading_ratio, 4),
            "total_body_words": total_body_words,
            "yellow_words": yellow_words,
            "green_words": green_words,
        },
        "qa_flags": [failure["type"] for failure in failures],
        "failures": failures,
        "warnings": warnings,
    }
    return report


def collect_docx(path: Path) -> list[Path]:
    if path.is_dir():
        return sorted(path.glob("*.docx"))
    return [path]


def main() -> int:
    parser = argparse.ArgumentParser(description="Lint Example Essay DOCX formatting and source highlighting.")
    parser.add_argument("paths", nargs="+", type=Path)
    parser.add_argument("--source-map", type=Path)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()

    reports = []
    for path in args.paths:
        for docx_path in collect_docx(path):
            try:
                reports.append(lint_docx(docx_path, args.source_map))
            except Exception as exc:
                reports.append({"docx": str(docx_path), "status": "fail", "qa_flags": ["docx_lint_error"], "error": str(exc)})

    result = {
        "status": "pass" if all(report.get("status") == "pass" for report in reports) else "fail",
        "reports": reports,
    }
    text = json.dumps(result, indent=2)
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(text, encoding="utf-8")
    else:
        print(text)
    return 0 if result["status"] == "pass" else 1


if __name__ == "__main__":
    raise SystemExit(main())
