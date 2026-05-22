#!/usr/bin/env python3
"""Read-only source inventory and text extraction for SBS exam workflow inputs."""

from __future__ import annotations

import argparse
import html
import json
import re
import sys
import zipfile
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

try:
    import fitz  # type: ignore
except Exception:  # pragma: no cover
    fitz = None

try:
    from pypdf import PdfReader  # type: ignore
except Exception:  # pragma: no cover
    PdfReader = None

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None

try:
    import openpyxl  # type: ignore
except Exception:  # pragma: no cover
    openpyxl = None

TEXT_EXTS = {".txt", ".md", ".markdown", ".yaml", ".yml", ".py", ".json", ".csv", ".tsv"}
SUPPORTED_EXTS = TEXT_EXTS | {".pdf", ".docx", ".pptx", ".pptm", ".ppsx", ".ppt", ".xlsx", ".xlsm"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".heic"}
COURSE_CODE_RE = re.compile(r"\b([A-Z]{4}\d{5}[A-Z]?)\b", re.IGNORECASE)
COURSE_CODE_FULL_RE = re.compile(r"[A-Z]{4}\d{5}[A-Z]?", re.IGNORECASE)


@dataclass
class ExtractedFile:
    path: str
    name: str
    extension: str
    role: str
    source_features: list[str]
    target_code: str | None
    target_name: str | None
    year: int | None
    source_trust_level: str
    evidence_use: str
    analysis_context: str
    allowed_factual_use: bool
    allowed_prediction_use: bool
    allowed_style_use: bool
    allowed_layout_use: bool
    allowed_transferable_example_use: bool
    status: str
    method: str
    text_path: str | None
    char_count: int
    page_count: int | None = None
    slide_count: int | None = None
    image_count: int | None = None
    limitations: list[str] = field(default_factory=list)
    preview: str = ""


def normalise(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def iter_files(paths: list[Path]) -> list[Path]:
    files: list[Path] = []
    for path in paths:
        if path.is_file():
            files.append(path)
        elif path.is_dir():
            files.extend(item for item in path.rglob("*") if item.is_file())
        else:
            files.append(path)
    return sorted(files, key=lambda item: str(item).lower())


def detect_year(path: Path, text: str, role: str | None = None) -> int | None:
    filename_matches = re.findall(r"(?<!\d)(20\d{2}|19\d{2})(?!\d)", path.name)
    if filename_matches:
        return int(filename_matches[-1])

    # Lecture decks often contain citation/reference years in slide text. Do
    # not treat those as source years; use metadata or explicit file names in
    # downstream code if a lecture-source year is required.
    if role in {"lecture_slide", "lecture_note", "annotated_lecture_slide", "essay_guidance"}:
        return None

    header = text[:3000]
    contextual = re.findall(
        r"(?:examination|exam|paper|assessment|semester|academic year|year)\D{0,80}\b(20\d{2}|19\d{2})\b",
        header,
        flags=re.IGNORECASE,
    )
    if contextual:
        return int(contextual[-1])
    return None


def detect_course_code(*values: str | None) -> str | None:
    for value in values:
        if not value:
            continue
        match = COURSE_CODE_RE.search(value)
        if match:
            return match.group(1).upper()
    return None


def clean_target_label(value: str | None) -> str | None:
    if not value:
        return None
    label = COURSE_CODE_RE.sub(" ", value)
    label = label.replace("_", " ").replace("-", " ")
    label = normalise(label)
    if not label or COURSE_CODE_FULL_RE.fullmatch(label):
        return None
    return label


def detect_target_group(
    path: Path,
    text: str,
    target_hint: str | None,
    target_label: str | None = None,
    target_code: str | None = None,
) -> tuple[str | None, str | None]:
    detected_code = detect_course_code(str(path), text[:5000])
    hint_code = detect_course_code(target_code, target_hint, target_label)
    hint_name = clean_target_label(target_label) or clean_target_label(target_hint)
    if detected_code:
        if hint_code and detected_code != hint_code:
            return detected_code, None
        return detected_code, hint_name
    if hint_code or hint_name:
        return hint_code, hint_name
    return None, None


def detect_source_features(path: Path, text: str) -> list[str]:
    name = path.name.lower()
    path_text = str(path).lower()
    body = (name + "\n" + text[:8000].lower())
    features: list[str] = []
    checks = [
        ("past_paper_path", "past paper" in path_text or "past papers" in path_text),
        ("formal_year_in_name", bool(re.search(r"(?<!\d)(20\d{2}|19\d{2})(?!\d)", name))),
        ("answer_key_or_solution", any(term in body for term in ("answer key", "with answer", "with answers", "answers only", "guide answer", "guide answers", "solutions:"))),
        ("example_or_mock_paper", any(term in body for term in ("example paper", "mock paper", "mock exam", "practice paper"))),
        ("practical_protocol", any(term in body for term in ("practical", "protocol", "workshop"))),
        ("essay_guidance_or_feedback", any(term in body for term in ("essay writing", "formative essay feedback", "essay feedback", "exam preparation", "exam guidance", "assessment guidance"))),
        ("marking_or_rubric", any(term in body for term in ("marking criteria", "rubric", "criteria for marking"))),
        ("mcq_or_single_best", any(term in body for term in ("mcq", "multiple choice", "single best answer"))),
        ("short_answer", "short answer" in body),
        ("essay_question", "essay" in body),
        ("problem_data_case", any(term in body for term in ("problem paper", "case study", "case studies", "data", "graph", "figure"))),
        ("lecture_or_slide", any(term in body for term in ("lecture", "slides", "learning objectives"))),
        ("recommended_reading", any(term in body for term in ("recommended reading", "reading list", "chapter", "textbook"))),
    ]
    for feature, present in checks:
        if present:
            features.append(feature)
    return features


def detect_role(path: Path, text: str, target_hint: str | None = None, source_features: list[str] | None = None) -> str:
    name = path.name.lower()
    path_text = str(path).lower()
    suffix = path.suffix.lower()
    body = (name + "\n" + text[:4000].lower())
    features = source_features or detect_source_features(path, text)
    if suffix in IMAGE_EXTS:
        if "lecture" in name or "slide" in name:
            return "annotated_lecture_slide"
        exemplar_terms = (
            "exemplar",
            "model",
            "answer",
            "essay",
            "draft",
            "example",
            "handwritten",
            "rewrite",
        )
        if any(term in body or term in path_text for term in exemplar_terms):
            return "exemplar_image"
        # Hashed WeChat/RWTemp images supplied alongside one target source set are
        # usually handwritten answer examples. Keep them style-only; factual
        # claims still require lecture/source verification.
        if target_hint and ("xwechat_files" in path_text or "rwtemp" in path_text):
            return "exemplar_image"
        return "unknown"
    if suffix in {".py", ".js", ".mjs", ".sh"}:
        return "helper_script"
    if suffix in {".yaml", ".yml", ".json"}:
        return "source_policy" if "policy" in name else "output_protocol" if "protocol" in name else "helper_script"
    if suffix in {".pptx", ".pptm", ".ppsx", ".ppt"}:
        if "essay_guidance_or_feedback" in features or any(term in name for term in ("essay writing", "tutorial", "exam guidance", "course information", "course info")):
            return "essay_guidance"
        if "answer_key_or_solution" in features:
            return "practice_answer_key"
        if any(term in name for term in ("quiz", "spotter", "formative", "practice", "mock", "problem sheet")):
            return "mock_exam" if "mock" in name else "practice_paper"
        return "lecture_slide"
    if "practical_protocol" in features and suffix == ".pdf":
        return "practical_protocol"
    if "example_or_mock_paper" in features and "past_paper_path" not in features:
        return "example_paper"
    if "past_paper_path" in features and "answer_key_or_solution" in features:
        return "formal_past_paper_with_answers"
    if "past_paper_path" in features and "formal_year_in_name" in features:
        return "formal_past_paper"
    if any(term in body for term in ("marking criteria", "rubric", "criteria for marking")):
        return "marking_criteria"
    if any(term in body for term in ("exemplar", "model answer", "example answer")):
        return "exemplar_answer"
    if "essay_guidance_or_feedback" in features or any(term in body for term in ("course information", "course info", "exam guidance", "assessment guidance")):
        return "essay_guidance"
    if "answer_key_or_solution" in features:
        return "practice_answer_key" if "practice" in body or "workshop" in body else "answer_key"
    if "example_or_mock_paper" in features:
        return "mock_exam" if "mock" in body else "example_paper"
    if any(term in body for term in ("quiz", "spotter", "formative", "practice question", "practice material", "problem sheet", "mcq", "multiple choice", "tutorial")) and not any(
        term in body for term in ("examination", "answer one question", "answer all questions", "section b")
    ):
        return "practice_paper"
    if any(term in body for term in ("examination", "answer one question", "answer all questions", "section a", "section b", "do not turn over")):
        return "formal_past_paper"
    if suffix == ".pdf" and any(term in body for term in ("lecture", "module", "learning objectives", "slide")):
        return "lecture_slide"
    if any(term in name for term in ("notes", "summary", "revision")):
        return "lecture_note"
    if "recommended_reading" in features:
        return "reading_list"
    return "unknown"


def trust_and_evidence(role: str) -> tuple[str, str]:
    if role in {"lecture_slide", "lecture_note", "annotated_lecture_slide"}:
        return "official_course", "factual_course_content"
    if role == "formal_past_paper":
        return "official_course", "formal_prediction_evidence"
    if role == "formal_past_paper_with_answers":
        return "official_course", "formal_prediction_and_answer_key_evidence"
    if role == "example_paper":
        return "official_course", "format_rule"
    if role == "practical_protocol":
        return "official_course", "practical_method_evidence"
    if role == "reading_list":
        return "official_course", "reading_recommendation"
    if role in {"practice_paper", "mock_exam"}:
        return "course_adjacent", "coverage_evidence_only"
    if role in {"answer_key", "practice_answer_key"}:
        return "course_adjacent", "answer_rationale_evidence"
    if role in {"exemplar_answer", "exemplar_image"}:
        return "course_adjacent", "answer_style_only"
    if role in {"marking_criteria", "essay_guidance"}:
        return "official_course", "format_rule"
    if role in {"source_policy", "output_protocol", "helper_script"}:
        return "course_adjacent", "excluded"
    return "unsupported" if role == "unsupported_binary" else "student_or_unknown", "excluded"


def _matches_target(
    source_code: str | None,
    source_name: str | None,
    target_label: str | None,
    target_code: str | None,
) -> bool:
    code = (source_code or "").lower()
    name = (source_name or "").lower()
    required_code = (target_code or "").lower()
    required_name = (target_label or "").lower()
    return bool(
        (required_code and code == required_code)
        or (required_name and code == required_name)
        or (required_name and (name == required_name or required_name in name or name in required_name))
    )


def infer_analysis_context(
    path: Path,
    role: str,
    target_group_code: str | None,
    target_group_name: str | None,
    target_label: str | None,
    required_target_code: str | None,
    example_mode: bool,
    status: str,
) -> str:
    path_text = str(path).lower()
    if status in {"failed", "unsupported"} or role in {"unknown", "unsupported_binary"}:
        return "unsupported_or_unreadable"
    if "benchmarks" in path_text or "fixture" in path_text or "fixtures" in path_text:
        return "benchmark_fixture"
    if role == "output_protocol":
        return "layout_exemplar"
    if role in {"exemplar_answer", "exemplar_image"}:
        return "style_exemplar"

    has_target = bool(target_label or required_target_code)
    same_target = _matches_target(target_group_code, target_group_name, target_label, required_target_code)
    if has_target and not same_target:
        return "cross_target_example" if example_mode or role in {
            "lecture_slide",
            "lecture_note",
            "annotated_lecture_slide",
            "formal_past_paper",
            "formal_past_paper_with_answers",
            "example_paper",
            "practice_paper",
            "practice_answer_key",
            "mock_exam",
            "answer_key",
            "practical_protocol",
            "reading_list",
            "marking_criteria",
            "essay_guidance",
        } else "unsupported_or_unreadable"
    if has_target and same_target:
        if role in {"lecture_slide", "lecture_note", "annotated_lecture_slide", "formal_past_paper", "formal_past_paper_with_answers", "marking_criteria"}:
            return "target_current_regime"
        return "target_auxiliary"

    if role in {"lecture_slide", "lecture_note", "annotated_lecture_slide", "formal_past_paper", "formal_past_paper_with_answers", "marking_criteria"}:
        return "target_current_regime"
    if role in {"practice_paper", "practice_answer_key", "mock_exam", "answer_key", "example_paper", "essay_guidance", "practical_protocol", "reading_list"}:
        return "target_auxiliary"
    return "unsupported_or_unreadable"


def allowed_use_flags(analysis_context: str, role: str) -> dict[str, bool]:
    style_roles = {"answer_key", "practice_answer_key", "exemplar_answer", "exemplar_image", "essay_guidance", "formal_past_paper_with_answers"}
    if analysis_context == "target_current_regime":
        return {
            "allowed_factual_use": role not in {"formal_past_paper", "formal_past_paper_with_answers"},
            "allowed_prediction_use": role in {"formal_past_paper", "formal_past_paper_with_answers"},
            "allowed_style_use": role in style_roles,
            "allowed_layout_use": False,
            "allowed_transferable_example_use": False,
        }
    if analysis_context == "target_old_or_different_regime":
        return {
            "allowed_factual_use": role in {"lecture_slide", "lecture_note", "annotated_lecture_slide", "marking_criteria"},
            "allowed_prediction_use": False,
            "allowed_style_use": False,
            "allowed_layout_use": False,
            "allowed_transferable_example_use": False,
        }
    if analysis_context == "target_auxiliary":
        return {
            "allowed_factual_use": role == "practical_protocol",
            "allowed_prediction_use": False,
            "allowed_style_use": role in style_roles,
            "allowed_layout_use": False,
            "allowed_transferable_example_use": False,
        }
    if analysis_context == "cross_target_example":
        return {
            "allowed_factual_use": False,
            "allowed_prediction_use": False,
            "allowed_style_use": False,
            "allowed_layout_use": False,
            "allowed_transferable_example_use": True,
        }
    if analysis_context == "style_exemplar":
        return {
            "allowed_factual_use": False,
            "allowed_prediction_use": False,
            "allowed_style_use": True,
            "allowed_layout_use": False,
            "allowed_transferable_example_use": False,
        }
    if analysis_context == "layout_exemplar":
        return {
            "allowed_factual_use": False,
            "allowed_prediction_use": False,
            "allowed_style_use": False,
            "allowed_layout_use": True,
            "allowed_transferable_example_use": False,
        }
    return {
        "allowed_factual_use": False,
        "allowed_prediction_use": False,
        "allowed_style_use": False,
        "allowed_layout_use": False,
        "allowed_transferable_example_use": False,
    }


def extract_pdf(path: Path) -> tuple[str, str, int | None, int | None, list[str]]:
    if fitz is not None:
        doc = fitz.open(path)
        pages = [page.get_text("text") for page in doc]
        images = sum(len(page.get_images(full=True)) for page in doc)
        limitations = ["PDF contains images; inspect rendered pages when image text or diagrams matter"] if images else []
        return "\n\n".join(pages), "pymupdf", doc.page_count, images, limitations
    if PdfReader is not None:
        reader = PdfReader(str(path))
        pages = [(page.extract_text() or "") for page in reader.pages]
        return "\n\n".join(pages), "pypdf", len(reader.pages), None, ["image count unavailable with pypdf fallback"]
    raise RuntimeError("no PDF extractor available")


def extract_docx(path: Path) -> tuple[str, str, list[str]]:
    if Document is None:
        raise RuntimeError("python-docx unavailable")
    doc = Document(str(path))
    parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts), "python-docx", []


def pptx_xml_text(xml_bytes: bytes) -> list[str]:
    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        raw = xml_bytes.decode("utf-8", errors="ignore")
        return [html.unescape(item) for item in re.findall(r"<a:t>(.*?)</a:t>", raw, flags=re.S)]
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    return [html.unescape(node.text or "") for node in root.findall(".//a:t", ns) if (node.text or "").strip()]


def extract_pptx(path: Path) -> tuple[str, str, int, list[str]]:
    parts: list[str] = []
    limitations = ["PPTX XML extraction captures text but not diagram/image-only content"]
    with zipfile.ZipFile(path) as zf:
        slide_names = sorted(
            (name for name in zf.namelist() if re.match(r"ppt/slides/slide\d+\.xml$", name)),
            key=lambda name: int(re.search(r"slide(\d+)\.xml", name).group(1)),  # type: ignore[union-attr]
        )
        for idx, name in enumerate(slide_names, start=1):
            texts = [normalise(text) for text in pptx_xml_text(zf.read(name))]
            texts = [text for text in texts if text]
            if texts:
                parts.append(f"SLIDE {idx}: " + " | ".join(texts))
        note_names = sorted(
            (name for name in zf.namelist() if re.match(r"ppt/notesSlides/notesSlide\d+\.xml$", name)),
            key=lambda name: int(re.search(r"notesSlide(\d+)\.xml", name).group(1)),  # type: ignore[union-attr]
        )
        for name in note_names:
            note_idx = int(re.search(r"notesSlide(\d+)\.xml", name).group(1))  # type: ignore[union-attr]
            texts = [normalise(text) for text in pptx_xml_text(zf.read(name))]
            texts = [text for text in texts if text and text.lower() != "slide"]
            if texts:
                parts.append(f"SLIDE {note_idx} NOTES: " + " | ".join(texts))
    return "\n".join(parts), "pptx-xml", len(slide_names), limitations


def extract_legacy_powerpoint(path: Path) -> tuple[str, str, list[str]]:
    data = path.read_bytes()
    ascii_chunks = re.findall(rb"[\x20-\x7E]{6,}", data)
    utf16_chunks = re.findall(rb"(?:[\x20-\x7E]\x00){6,}", data)
    parts = [chunk.decode("utf-8", errors="ignore") for chunk in ascii_chunks]
    parts.extend(chunk.decode("utf-16le", errors="ignore") for chunk in utf16_chunks)
    seen = set()
    cleaned: list[str] = []
    for part in parts:
        text = normalise(part)
        if text and text not in seen:
            cleaned.append(text)
            seen.add(text)
    return "\n".join(cleaned[:2000]), "binary-strings", [
        "Legacy .ppt extraction is approximate; inspect the original file when diagrams or exact wording matter"
    ]


def extract_xlsx(path: Path) -> tuple[str, str, list[str]]:
    if openpyxl is None:
        raise RuntimeError("openpyxl unavailable")
    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"SHEET: {sheet.title}")
        row_count = 0
        for row in sheet.iter_rows(values_only=True):
            values = [normalise(str(value)) for value in row if value is not None and normalise(str(value))]
            if values:
                parts.append(" | ".join(values))
                row_count += 1
            if row_count >= 500:
                parts.append("[sheet truncated after 500 non-empty rows]")
                break
    return "\n".join(parts), "openpyxl", []


def extract_text(path: Path) -> tuple[str, str, dict[str, Any], list[str]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text, method, page_count, image_count, limitations = extract_pdf(path)
        return text, method, {"page_count": page_count, "image_count": image_count}, limitations
    if suffix == ".docx":
        text, method, limitations = extract_docx(path)
        return text, method, {}, limitations
    if suffix in {".pptx", ".pptm", ".ppsx"}:
        text, method, slide_count, limitations = extract_pptx(path)
        return text, method, {"slide_count": slide_count}, limitations
    if suffix == ".ppt":
        text, method, limitations = extract_legacy_powerpoint(path)
        return text, method, {}, limitations
    if suffix in {".xlsx", ".xlsm"}:
        text, method, limitations = extract_xlsx(path)
        return text, method, {}, limitations
    if suffix in TEXT_EXTS:
        return path.read_text(encoding="utf-8", errors="replace"), "text", {}, []
    if suffix in IMAGE_EXTS:
        return "", "image-metadata", {}, ["OCR not performed by this script"]
    raise RuntimeError(f"unsupported extension: {suffix}")


def safe_text_name(path: Path, index: int) -> str:
    base = re.sub(r"[^A-Za-z0-9_.-]+", "_", path.stem).strip("_")[:90] or "source"
    return f"{index:03d}_{base}.txt"


def scan(
    paths: list[Path],
    output_dir: Path,
    target_hint: str | None,
    target_label: str | None = None,
    target_code: str | None = None,
    example_mode: bool = False,
    benchmark_fixture_config: Path | None = None,
) -> dict[str, Any]:
    text_dir = output_dir / "source_text"
    text_dir.mkdir(parents=True, exist_ok=True)
    rows: list[ExtractedFile] = []
    for index, path in enumerate(iter_files(paths), start=1):
        suffix = path.suffix.lower()
        status = "unsupported"
        method = "none"
        text = ""
        extra: dict[str, Any] = {}
        limitations: list[str] = []
        text_path: str | None = None
        if not path.exists():
            status = "failed"
            limitations = ["path does not exist"]
        elif suffix not in SUPPORTED_EXTS and suffix not in IMAGE_EXTS:
            limitations = [f"unsupported extension: {suffix or '[none]'}"]
            if suffix in {".pyc", ".bin", ".exe"}:
                status = "unsupported"
        else:
            try:
                text, method, extra, limitations = extract_text(path)
                status = "partial" if suffix in IMAGE_EXTS else "ok"
                if text.strip():
                    out = text_dir / safe_text_name(path, index)
                    out.write_text(text, encoding="utf-8")
                    text_path = str(out)
            except Exception as exc:  # pragma: no cover
                status = "failed"
                method = method if method != "none" else suffix.lstrip(".") or "unknown"
                limitations = [f"{type(exc).__name__}: {exc}"]
        detected_target_code, detected_target_name = detect_target_group(
            path,
            text,
            target_hint,
            target_label=target_label,
            target_code=target_code,
        )
        source_features = detect_source_features(path, text)
        role = detect_role(path, text, target_hint, source_features)
        if status == "unsupported":
            role = "unsupported_binary"
        source_trust_level, evidence_use = trust_and_evidence(role)
        analysis_context = infer_analysis_context(
            path=path,
            role=role,
            target_group_code=detected_target_code,
            target_group_name=detected_target_name,
            target_label=target_label or target_hint,
            required_target_code=target_code,
            example_mode=example_mode,
            status=status,
        )
        allowed_flags = allowed_use_flags(analysis_context, role)
        if role == "exemplar_image":
            if "visual_inspection_required" not in limitations:
                limitations.append("visual_inspection_required")
            limitations.append(
                "Image exemplar may be used for essay style, paragraph logic, and density only; factual claims require verification from lecture or reliable sources."
            )
        year = detect_year(path, text, role)
        rows.append(
            ExtractedFile(
                path=str(path),
                name=path.name,
                extension=suffix,
                role=role,
                source_features=source_features,
                target_code=detected_target_code,
                target_name=detected_target_name,
                year=year,
                source_trust_level=source_trust_level,
                evidence_use=evidence_use,
                analysis_context=analysis_context,
                allowed_factual_use=allowed_flags["allowed_factual_use"],
                allowed_prediction_use=allowed_flags["allowed_prediction_use"],
                allowed_style_use=allowed_flags["allowed_style_use"],
                allowed_layout_use=allowed_flags["allowed_layout_use"],
                allowed_transferable_example_use=allowed_flags["allowed_transferable_example_use"],
                status=status,
                method=method,
                text_path=text_path,
                char_count=len(text),
                page_count=extra.get("page_count"),
                slide_count=extra.get("slide_count"),
                image_count=extra.get("image_count"),
                limitations=limitations,
                preview=normalise(text[:500]),
            )
        )
    manifest = {
        "inputs": [str(path) for path in paths],
        "target_hint": target_hint,
        "target_label": target_label,
        "target_code": target_code,
        "example_mode": example_mode,
        "benchmark_fixture_config": str(benchmark_fixture_config) if benchmark_fixture_config else None,
        "file_count": len(rows),
        "files": [asdict(row) for row in rows],
    }
    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "source_scan.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("paths", nargs="+", help="Files or directories to scan")
    parser.add_argument("--output", required=True, help="Output directory for manifest and source_text")
    parser.add_argument("--target-hint", default=None, help="Optional target group hint")
    parser.add_argument("--target", default=None, help="Target course/module/source-set label")
    parser.add_argument("--target-code", default=None, help="Target course/module code")
    parser.add_argument("--example-mode", action="store_true", help="Classify non-target sources as transferable examples")
    parser.add_argument("--benchmark-fixture-config", type=Path, default=None, help="Optional benchmark fixture config path")
    args = parser.parse_args(argv)
    manifest = scan(
        [Path(path).expanduser() for path in args.paths],
        Path(args.output).expanduser(),
        args.target_hint,
        target_label=args.target,
        target_code=args.target_code,
        example_mode=args.example_mode,
        benchmark_fixture_config=args.benchmark_fixture_config,
    )
    counts: dict[str, int] = {}
    for row in manifest["files"]:
        counts[row["role"]] = counts.get(row["role"], 0) + 1
    print(json.dumps({"file_count": manifest["file_count"], "role_counts": counts}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
