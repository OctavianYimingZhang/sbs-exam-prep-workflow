#!/usr/bin/env python3
"""Lint Academic Exam-Ready Notes DOCX style."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Cm, RGBColor  # type: ignore
except Exception as exc:  # pragma: no cover
    raise SystemExit(f"python-docx is required: {exc}")


BLUE_RGB = {"0000FF", "0563C1", "2F5496", "1F4E79"}
EXPECTED_MARGIN_CM = 2.0
MAX_LINE_SPACING = 1.2
MIN_LINE_SPACING = 1.0
FORBIDDEN_INTERNAL_HEADINGS = {
    "Exam Specificity",
    "Core Exam Claim",
    "Exam Use",
    "Common Error / Trap",
    "Must Master",
    "Course-Level Exam Map",
    "How To Answer This Exam",
}


def iter_docx_paths(path: Path) -> list[Path]:
    if path.is_file() and path.suffix.lower() == ".docx":
        return [path]
    if path.is_dir():
        return sorted(child for child in path.rglob("*.docx") if child.is_file())
    return []


def cm_value(length: Any) -> float:
    return float(length.cm)


def line_spacing_value(paragraph: Any) -> float | None:
    value = paragraph.paragraph_format.line_spacing
    if value is None and paragraph.style is not None:
        value = paragraph.style.paragraph_format.line_spacing
    if value is None:
        return None
    try:
        return float(value)
    except TypeError:
        return None


def count_page_breaks(path: Path) -> int:
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
    explicit_breaks = len(re.findall(r"<w:br\b[^>]*w:type=[\"']page[\"']", xml))
    page_break_properties = len(re.findall(r"<w:pageBreakBefore\b", xml))
    return explicit_breaks + page_break_properties


def is_forbidden_heading(text: str) -> bool:
    normalized = text.strip().rstrip(":").casefold()
    return normalized in {heading.casefold() for heading in FORBIDDEN_INTERNAL_HEADINGS}


def is_body_paragraph(paragraph: Any) -> bool:
    style_name = (paragraph.style.name if paragraph.style else "").casefold()
    text = paragraph.text.strip()
    return bool(text) and not style_name.startswith("heading") and style_name not in {"epntitle", "epnlecture", "title"}


def lint_docx(path: Path) -> list[dict[str, Any]]:
    failures: list[dict[str, Any]] = []
    doc = Document(path)
    if not doc.sections:
        failures.append({"type": "missing_docx_section", "path": str(path)})
        return failures

    section = doc.sections[0]
    for margin_name in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        value = cm_value(getattr(section, margin_name))
        if abs(value - EXPECTED_MARGIN_CM) > 0.08:
            failures.append({"type": "bad_margin", "path": str(path), "margin": margin_name, "cm": round(value, 3)})

    lecture_headings = 0
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        if is_forbidden_heading(text):
            failures.append({"type": "forbidden_internal_heading", "path": str(path), "paragraph": index, "text": text})
        if text.casefold().startswith("lecture:"):
            lecture_headings += 1
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            if paragraph.alignment not in {None, WD_ALIGN_PARAGRAPH.LEFT}:
                failures.append({"type": "heading_not_left_aligned", "path": str(path), "paragraph": index})
        spacing = line_spacing_value(paragraph)
        if spacing is not None and not (MIN_LINE_SPACING <= spacing <= MAX_LINE_SPACING):
            failures.append({"type": "bad_line_spacing", "path": str(path), "paragraph": index, "line_spacing": round(spacing, 3)})
        if is_body_paragraph(paragraph) and paragraph.alignment not in {None, WD_ALIGN_PARAGRAPH.LEFT}:
            failures.append({"type": "body_not_left_aligned", "path": str(path), "paragraph": index, "alignment": str(paragraph.alignment)})
        if is_body_paragraph(paragraph) and paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            failures.append({"type": "body_justified_like_essay", "path": str(path), "paragraph": index})
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
            if color is not None and str(color).upper() != "000000":
                failures.append({"type": "non_black_text", "path": str(path), "paragraph": index, "text": run.text[:80], "rgb": str(color)})
            if color is not None and str(color).upper() in BLUE_RGB:
                failures.append({"type": "blue_text_detected", "path": str(path), "paragraph": index, "text": run.text[:80]})
            font_names = {name for name in [run.font.name, paragraph.style.font.name if paragraph.style else None] if name}
            if font_names and "Arial" not in font_names:
                failures.append({"type": "non_arial_text", "path": str(path), "paragraph": index, "fonts": sorted(font_names)})

    page_breaks = count_page_breaks(path)
    if lecture_headings and page_breaks < max(0, lecture_headings - 1):
        failures.append({"type": "missing_lecture_page_breaks", "path": str(path), "lecture_headings": lecture_headings, "page_breaks": page_breaks})

    return failures


def lint_path(path: Path) -> dict[str, Any]:
    docx_paths = iter_docx_paths(path)
    failures: list[dict[str, Any]] = []
    if not docx_paths:
        failures.append({"type": "no_docx_found", "path": str(path)})
    for docx_path in docx_paths:
        failures.extend(lint_docx(docx_path))
    return {
        "pass": not failures,
        "counts": {"docx_files": len(docx_paths)},
        "failures": failures,
    }


def create_bad_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    paragraph = doc.add_paragraph("Exam Use:")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run("This blue text must fail.")
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 0, 255)
    doc.save(path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("path", type=Path, nargs="?")
    parser.add_argument("--self-test-bad", action="store_true")
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()

    tmp_dir: str | None = None
    try:
        if args.self_test_bad:
            tmp_dir = tempfile.mkdtemp(prefix="exam_prep_style_bad_")
            target = Path(tmp_dir) / "bad_style.docx"
            create_bad_docx(target)
            result = lint_path(target)
        elif args.path:
            result = lint_path(args.path)
        else:
            result = {"pass": False, "failures": [{"type": "missing_path"}], "counts": {"docx_files": 0}}
    except Exception as exc:
        result = {"pass": False, "failures": [{"type": "read_error", "error": str(exc)}], "counts": {"docx_files": 0}}
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)

    text = json.dumps(result, indent=2)
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(text + "\n", encoding="utf-8")
    else:
        print(text)
    return 0 if result.get("pass") else 1


if __name__ == "__main__":
    raise SystemExit(main())
