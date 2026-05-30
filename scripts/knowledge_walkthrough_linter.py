#!/usr/bin/env python3
"""Lint Knowledge Walkthrough DOCX files for formatting and student-output policy."""

from __future__ import annotations

import argparse
import json
import re
import tempfile
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Cm, Pt  # type: ignore
except Exception as exc:  # pragma: no cover
    raise SystemExit(f"python-docx is required: {exc}")

from knowledge_only_rendering_rules import (
    forbidden_advisory_heading_hits,
    forbidden_advisory_phrase_hits,
    forbidden_non_knowledge_hits,
    repeated_template_label_hits,
)


STYLE_LIMITS = {
    "margin_cm": 2.0,
    "margin_tolerance_cm": 0.08,
    "line_spacing_min": 1.0,
    "line_spacing_max": 1.2,
    "body_font_min_pt": 9.5,
    "body_font_max_pt": 11.5,
}


FORBIDDEN_TEXT = {
    "source anchor",
    "confidence",
    "evidence score",
    "recurrence count",
    "examiner operation",
    "discriminator axis",
    "essay plan",
    "full example essay",
    "practice question",
    "answer key",
    "past paper year",
    "prediction score",
    "according to slide",
    "english explanations extracted",
    "ppt page",
    "slide mentions",
    "slides say",
    "the final slide",
    "the first slide",
    "the next slide",
    "the second slide",
    "this slide",
}

OLD_VISIBLE_SCAFFOLD_HEADINGS = {
    "How To Answer This Exam",
    "What This Lecture Is About",
    "What This Module Explains",
    "Knowledge Walkthrough",
    "Key Logic",
    "Knowledge Points",
    "Must Master",
    "Lecture Recap",
}

SEMANTIC_HEADINGS = {"Knowledge map", "Key distinctions", "Core knowledge points", "Synthesis"}


def collect_docx(path: Path) -> list[Path]:
    if path.is_dir():
        return sorted(path.glob("*.docx"))
    return [path]


def run_size_pt(run) -> float | None:
    if run.font.size is None:
        return None
    return float(run.font.size.pt)


def build_bad_style_fixture(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.5
    title.add_run("Lecture Knowledge Walkthrough").font.name = "Times New Roman"
    for text in [
        "English explanations extracted from the shared ChatGPT page",
        "This slide shows the old source-route wrapper.",
        "Lecture: Fixture Lecture",
        "What This Lecture Is About",
        "Module Map",
        "Knowledge Walkthrough",
        "Key Logic",
        "Must Master",
        "Definition: one",
        "Definition: two",
        "Definition: three",
        "Definition: four",
    ]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    doc.save(path)


def lint_docx(path: Path) -> dict[str, Any]:
    doc = Document(path)
    failures: list[dict[str, Any]] = []
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    lowered = text.lower()
    for phrase in sorted(FORBIDDEN_TEXT):
        if phrase in lowered:
            failures.append({"type": "forbidden_student_text", "phrase": phrase})
    for phrase in forbidden_advisory_phrase_hits(text):
        failures.append({"type": "forbidden_advisory_phrase", "phrase": phrase})
    for heading in forbidden_advisory_heading_hits(text):
        failures.append({"type": "forbidden_advisory_heading", "heading": heading})
    for category in forbidden_non_knowledge_hits(text):
        failures.append({"type": "forbidden_non_knowledge_surface", "category": category})
    for label in repeated_template_label_hits(text):
        failures.append({"type": "repeated_rigid_template_label", "label": label})
    for heading in OLD_VISIBLE_SCAFFOLD_HEADINGS:
        if re.search(rf"(?im)^\s*{re.escape(heading)}\s*:?\s*$", text):
            failures.append({"type": "old_visible_scaffold_heading", "heading": heading})

    section = doc.sections[0]
    margins_cm = [section.top_margin.cm, section.bottom_margin.cm, section.left_margin.cm, section.right_margin.cm]
    if any(abs(value - STYLE_LIMITS["margin_cm"]) > STYLE_LIMITS["margin_tolerance_cm"] for value in margins_cm):
        failures.append({"type": "margins_not_compact_2_0_cm", "margins_cm": margins_cm})

    heading_terms = set(SEMANTIC_HEADINGS)
    visible_index = 0
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
        visible_index += 1
        line_spacing = paragraph.paragraph_format.line_spacing
        if line_spacing is None or not (STYLE_LIMITS["line_spacing_min"] <= float(line_spacing) <= STYLE_LIMITS["line_spacing_max"]):
            failures.append({"type": "line_spacing_not_compact", "paragraph": visible_index, "line_spacing": line_spacing})
        if visible_index == 1:
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                failures.append({"type": "title_not_centered"})
        elif paragraph.text in heading_terms or paragraph.text.startswith("Lecture:"):
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                failures.append({"type": "heading_not_left", "paragraph": visible_index, "text": paragraph.text[:80]})
        else:
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                failures.append({"type": "body_not_left_aligned", "paragraph": visible_index, "text": paragraph.text[:80]})
        for run_idx, run in enumerate([run for run in paragraph.runs if run.text], start=1):
            if run.font.name != "Arial":
                failures.append({"type": "run_font_not_arial", "paragraph": visible_index, "run": run_idx, "font": run.font.name})
            size = run_size_pt(run)
            if size is not None and not (STYLE_LIMITS["body_font_min_pt"] <= size <= 15.0):
                failures.append({"type": "run_font_size_outside_compact_range", "paragraph": visible_index, "run": run_idx, "font_size_pt": size})
            color = run.font.color.rgb
            if color is not None and str(color) != "000000":
                failures.append({"type": "run_font_color_not_black", "paragraph": visible_index, "run": run_idx, "color": str(color)})

    return {
        "docx": str(path),
        "status": "pass" if not failures else "fail",
        "failures": failures,
        "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("paths", nargs="*", type=Path)
    parser.add_argument("--output", type=Path)
    parser.add_argument("--self-test-bad", action="store_true")
    args = parser.parse_args()
    if args.self_test_bad:
        with tempfile.TemporaryDirectory(prefix="kw_bad_style_") as tmp:
            bad_path = Path(tmp) / "bad_knowledge_walkthrough.docx"
            build_bad_style_fixture(bad_path)
            result = {"status": "fail", "reports": [lint_docx(bad_path)]}
            result["status"] = "pass" if all(report["status"] == "pass" for report in result["reports"]) else "fail"
            print(json.dumps(result, indent=2))
            return 0 if result["status"] == "pass" else 1
    reports = []
    for path in args.paths:
        for docx in collect_docx(path):
            reports.append(lint_docx(docx))
    if not reports:
        reports.append({"status": "fail", "failures": [{"type": "no_docx_found"}]})
    result = {"status": "pass" if all(report["status"] == "pass" for report in reports) else "fail", "reports": reports}
    text = json.dumps(result, indent=2)
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(text, encoding="utf-8")
    else:
        print(text)
    return 0 if result["status"] == "pass" else 1


if __name__ == "__main__":
    raise SystemExit(main())
