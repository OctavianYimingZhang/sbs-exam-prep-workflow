#!/usr/bin/env python3
"""Generate standalone DOCX files for explicit Example Essay Mode.

Input is a JSON document containing one or more ExampleEssayDocumentPlan-like
records. The script writes one DOCX per essay plus internal QA/source-map
artefacts. It enforces the formatting contract directly on every paragraph and
run.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import zipfile
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Cm, Pt  # type: ignore
except Exception as exc:  # pragma: no cover
    raise SystemExit(f"python-docx is required: {exc}")


BODY_KIND = {"body", "paragraph", "conclusion", "intro", "introduction", "mechanism", "evidence", "evaluation", "synthesis"}
AUTHOR_YEAR_RE = re.compile(r"\([A-Z][A-Za-z'’-]+(?:\s+et\s+al\.|\s+and\s+[A-Z][A-Za-z'’-]+)?(?:,\s*|\s+)(?:19|20)\d{2}[a-z]?\)")
MICRO_DETAIL_SOURCE_TYPES = {"extra_reading_book", "citation_original_source", "classic_experiment_source"}
MICRO_DETAIL_ALLOWED_CLAIM_DELTAS = {None, "", "precision_only", "none"}


def safe_filename(text: str, max_len: int = 72) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]+", "", text)
    cleaned = re.sub(r"\s+", "_", cleaned.strip())
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "", cleaned)
    return (cleaned[:max_len].strip("._-") or "example_essay")


def word_count(text: str) -> int:
    return len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?", text))


def as_float(value: Any) -> float | None:
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        try:
            return float(value.strip().rstrip("%"))
        except ValueError:
            return None
    return None


def paragraph_kind(paragraph: dict[str, Any]) -> str:
    if paragraph.get("is_title"):
        return "title"
    if paragraph.get("is_subtitle"):
        return "subtitle"
    if paragraph.get("is_heading"):
        return "heading"
    return "body"


def normalize_paragraph(paragraph, kind: str) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.5
    if kind == "title":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif kind in {"subtitle", "heading"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        normalize_run(run)


def normalize_run(run) -> None:
    run.font.name = "Arial"
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.paragraph_format.line_spacing = 1.5

    for style_name in ["EssayTitle", "EssaySubtitle", "EssayHeading", "EssayBody"]:
        if style_name not in styles:
            styles.add_style(style_name, 1)
        style = styles[style_name]
        style.font.name = "Arial"
        style.paragraph_format.line_spacing = 1.5


def add_text_run(paragraph, run_data: dict[str, Any]) -> None:
    run = paragraph.add_run(str(run_data.get("text", "")))
    normalize_run(run)
    highlight = run_data.get("highlight", "none")
    if highlight == "yellow":
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    elif highlight == "green":
        run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN


def validate_plan(essay: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    if not essay.get("lecture_anchors"):
        errors.append("essay_requires_lecture_anchors")
    compression_budget = essay.get("compression_budget")
    if isinstance(compression_budget, dict):
        protected_source_skeleton = compression_budget.get("protected_source_skeleton") or []
        if not protected_source_skeleton:
            errors.append("compression_budget_missing_protected_source_skeleton")
        requested = compression_budget.get("requested_reduction") or {}
        requested_type = requested.get("type") if isinstance(requested, dict) else None
        requested_value = as_float(requested.get("value")) if isinstance(requested, dict) else None
        safe_range = compression_budget.get("safe_reduction_range") or {}
        safe_max = as_float(safe_range.get("max")) if isinstance(safe_range, dict) else None
        decision = compression_budget.get("decision")
        if requested_type == "percent" and requested_value is not None and safe_max is not None and requested_value > safe_max:
            if decision != "reject_requested_reduction":
                errors.append("compression_target_exceeds_safe_budget")
    body_paragraphs = [p for p in essay.get("paragraphs", []) if paragraph_kind(p) == "body"]
    if not body_paragraphs:
        errors.append("essay_requires_body_paragraphs")
    for idx, paragraph in enumerate(essay.get("paragraphs", []), start=1):
        kind = paragraph_kind(paragraph)
        if kind == "body" and not paragraph.get("lecture_anchors"):
            errors.append(f"paragraph_{idx}_missing_lecture_anchor")
        for run_idx, run in enumerate(paragraph.get("text_runs", []), start=1):
            source_type = run.get("source_type")
            highlight = run.get("highlight", "none")
            if source_type in {"citation_original_source", "classic_experiment_source"}:
                if highlight != "green":
                    errors.append(f"paragraph_{idx}_run_{run_idx}_{source_type}_requires_green")
                if not run.get("in_text_citation") and not AUTHOR_YEAR_RE.search(str(run.get("text", ""))):
                    errors.append(f"paragraph_{idx}_run_{run_idx}_green_run_missing_author_year")
                if run.get("citation_original_read") is not True:
                    errors.append(f"paragraph_{idx}_run_{run_idx}_{source_type}_not_read")
            if source_type == "extra_reading_book":
                if highlight != "yellow":
                    errors.append(f"paragraph_{idx}_run_{run_idx}_extra_reading_requires_yellow")
                if not run.get("source_anchor"):
                    errors.append(f"paragraph_{idx}_run_{run_idx}_extra_reading_missing_chapter_anchor")
            if highlight == "green" and source_type not in {"citation_original_source", "classic_experiment_source"}:
                errors.append(f"paragraph_{idx}_run_{run_idx}_green_wrong_source_type")
            if highlight == "yellow" and source_type != "extra_reading_book":
                errors.append(f"paragraph_{idx}_run_{run_idx}_yellow_wrong_source_type")
            if run.get("micro_detail_insert") is True:
                inserted_phrase = str(run.get("inserted_phrase", "")).strip()
                if source_type not in MICRO_DETAIL_SOURCE_TYPES:
                    errors.append(f"paragraph_{idx}_run_{run_idx}_micro_detail_requires_verified_extra_source")
                if not run.get("parent_ppt_or_source_slot"):
                    errors.append(f"paragraph_{idx}_run_{run_idx}_micro_detail_parent_slot_missing")
                if not run.get("source_anchor"):
                    errors.append(f"paragraph_{idx}_run_{run_idx}_micro_detail_insert_missing_source_anchor")
                if not inserted_phrase:
                    errors.append(f"paragraph_{idx}_run_{run_idx}_micro_detail_missing_inserted_phrase")
                if run.get("claim_delta") not in MICRO_DETAIL_ALLOWED_CLAIM_DELTAS:
                    errors.append(f"paragraph_{idx}_run_{run_idx}_micro_detail_claim_delta_not_precision_only")
    return errors


def write_essay(essay: dict[str, Any], out_dir: Path, qa_dir: Path, index: int) -> dict[str, Any]:
    qa_flags = list(essay.get("qa_flags", []))
    validation_errors = validate_plan(essay)
    qa_flags.extend(validation_errors)

    essay_id = essay.get("essay_id") or f"EE{index:02d}"
    title = essay.get("title") or essay.get("question") or f"Example Essay {index}"
    filename = f"{essay_id}_{safe_filename(title)}.docx"
    docx_path = out_dir / filename

    doc = Document()
    set_document_defaults(doc)

    source_map = {
        "essay_id": essay_id,
        "question": essay.get("question"),
        "docx_filename": filename,
        "target_group_key": essay.get("target_group_key"),
        "lecture_anchors": essay.get("lecture_anchors", []),
        "extra_reading_status": essay.get("extra_reading_status", "not_supplied"),
        "compression_budget": essay.get("compression_budget"),
        "paragraphs": [],
        "qa_flags": qa_flags,
    }

    total_body_words = 0
    yellow_words = 0
    green_words = 0
    micro_detail_enhancements: list[dict[str, Any]] = []

    for p_idx, paragraph_data in enumerate(essay.get("paragraphs", []), start=1):
        kind = paragraph_kind(paragraph_data)
        style = {
            "title": "EssayTitle",
            "subtitle": "EssaySubtitle",
            "heading": "EssayHeading",
            "body": "EssayBody",
        }[kind]
        paragraph = doc.add_paragraph(style=style)
        if paragraph_data.get("is_heading"):
            paragraph.paragraph_format.keep_with_next = True
        if kind in {"title", "heading"}:
            # Bold is allowed for title/section headings.
            pass

        para_map = {
            "paragraph_index": p_idx,
            "paragraph_id": paragraph_data.get("paragraph_id", f"P{p_idx:02d}"),
            "kind": kind,
            "function": paragraph_data.get("function"),
            "lecture_anchors": paragraph_data.get("lecture_anchors", []),
            "runs": [],
        }

        for r_idx, run_data in enumerate(paragraph_data.get("text_runs", []), start=1):
            add_text_run(paragraph, run_data)
            if kind in BODY_KIND or kind == "body":
                wc = word_count(str(run_data.get("text", "")))
                total_body_words += wc
                if run_data.get("highlight") == "yellow":
                    yellow_words += wc
                if run_data.get("highlight") == "green":
                    green_words += wc
            run_map = {
                "run_index": r_idx,
                "source_type": run_data.get("source_type"),
                "source_anchor": run_data.get("source_anchor"),
                "highlight": run_data.get("highlight", "none"),
                "in_text_citation": run_data.get("in_text_citation"),
                "citation_original_read": run_data.get("citation_original_read"),
                "word_count": word_count(str(run_data.get("text", ""))),
                "micro_detail_insert": run_data.get("micro_detail_insert", False),
                "original_phrase": run_data.get("original_phrase"),
                "inserted_phrase": run_data.get("inserted_phrase"),
                "parent_ppt_or_source_slot": run_data.get("parent_ppt_or_source_slot"),
                "question_function": run_data.get("question_function"),
                "claim_delta": run_data.get("claim_delta"),
                "qa_status": run_data.get("qa_status"),
            }
            para_map["runs"].append(run_map)
            if run_map["micro_detail_insert"]:
                micro_detail_enhancements.append(
                    {
                        "paragraph_id": para_map["paragraph_id"],
                        **run_map,
                    }
                )
        if kind in {"title", "heading"}:
            for run in paragraph.runs:
                run.bold = True
                normalize_run(run)
        normalize_paragraph(paragraph, kind)
        source_map["paragraphs"].append(para_map)

    source_map["word_counts"] = {
        "total_body_words": total_body_words,
        "extra_reading_yellow_words": yellow_words,
        "citation_original_green_words": green_words,
        "extra_reading_ratio": (yellow_words / total_body_words) if total_body_words else 0.0,
    }
    source_map["micro_detail_enhancements"] = micro_detail_enhancements

    doc.save(docx_path)

    source_map_path = qa_dir / f"{essay_id}_source_map.json"
    qa_path = qa_dir / f"{essay_id}_qa.json"
    source_map_path.write_text(json.dumps(source_map, indent=2), encoding="utf-8")
    qa_path.write_text(json.dumps({"essay_id": essay_id, "qa_flags": qa_flags, "validation_errors": validation_errors}, indent=2), encoding="utf-8")

    return {
        "essay_id": essay_id,
        "question": essay.get("question"),
        "title": title,
        "docx_filename": filename,
        "docx_path": str(docx_path),
        "source_map": str(source_map_path),
        "qa": str(qa_path),
        "qa_flags": qa_flags,
        "word_counts": source_map["word_counts"],
        "micro_detail_enhancements": micro_detail_enhancements,
    }


def load_essays(path: Path) -> list[dict[str, Any]]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if isinstance(data, list):
        return data
    if isinstance(data.get("essays"), list):
        return data["essays"]
    return [data]


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate standalone DOCX files for Example Essay Mode.")
    parser.add_argument("--plan", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--zip", action="store_true", dest="make_zip")
    parser.add_argument("--clean", action="store_true")
    parser.add_argument("--strict", action="store_true", help="Fail if any generated essay records validation QA flags")
    parser.add_argument("--qa-dir", type=Path, help="Optional internal QA artefact directory. Use this to keep the public output directory deliverable-only.")
    parser.add_argument("--deliverable-only", action="store_true", help="Write helper JSON outside the public output directory.")
    args = parser.parse_args()

    if args.clean and args.output_dir.exists():
        shutil.rmtree(args.output_dir)
    args.output_dir.mkdir(parents=True, exist_ok=True)
    qa_dir = args.qa_dir or (args.output_dir.parent / f"{args.output_dir.name}_internal_qa" if args.deliverable_only else args.output_dir)
    if args.clean and qa_dir.exists() and qa_dir != args.output_dir:
        shutil.rmtree(qa_dir)
    qa_dir.mkdir(parents=True, exist_ok=True)

    essays = load_essays(args.plan)
    manifest = {
        "plan": str(args.plan),
        "output_dir": str(args.output_dir),
        "documents": [],
    }
    for idx, essay in enumerate(essays, start=1):
        manifest["documents"].append(write_essay(essay, args.output_dir, qa_dir, idx))

    manifest_path = qa_dir / "example_essay_manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    combined_audit = {
        "essays": [
            {
                "essay_id": doc["essay_id"],
                "question": doc["question"],
                "docx_filename": doc["docx_filename"],
                "source_map": doc["source_map"],
                "qa_flags": doc["qa_flags"],
                "word_counts": doc["word_counts"],
                "micro_detail_enhancements": doc.get("micro_detail_enhancements", []),
            }
            for doc in manifest["documents"]
        ]
    }
    audit_path = qa_dir / "example_essay_source_audit.json"
    audit_path.write_text(json.dumps(combined_audit, indent=2), encoding="utf-8")

    if args.make_zip:
        zip_path = args.output_dir.with_suffix(".zip")
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for item in args.output_dir.rglob("*"):
                if item.is_file():
                    zf.write(item, item.relative_to(args.output_dir.parent))
        manifest["zip"] = str(zip_path)
        manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    strict_failures = [
        {"essay_id": doc["essay_id"], "qa_flags": doc["qa_flags"]}
        for doc in manifest["documents"]
        if doc.get("qa_flags")
    ]
    status = "fail" if args.strict and strict_failures else "ok"
    print(
        json.dumps(
            {
                "status": status,
                "manifest": str(manifest_path),
                "audit": str(audit_path),
                "documents": manifest["documents"],
                "strict_failures": strict_failures,
            },
            indent=2,
        )
    )
    return 1 if status == "fail" else 0


if __name__ == "__main__":
    raise SystemExit(main())
