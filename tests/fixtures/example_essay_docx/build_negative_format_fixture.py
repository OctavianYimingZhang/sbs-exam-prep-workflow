#!/usr/bin/env python3
"""Build a DOCX fixture that should fail docx_format_linter.py."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def main() -> None:
    out = Path(__file__).with_name("negative_format_fixture.docx")
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("This fixture intentionally uses wrong margins, Times New Roman, paragraph spacing and left-aligned body text.")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
